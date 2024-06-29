import pypdf
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

def extract_toc_and_content(pdf_path):
    reader = pypdf.PdfReader(pdf_path)
    toc = reader.outline
    content = ""
    for page in reader.pages:
        content += page.extract_text()
    return toc, content

def process_toc_item(item, level=0):
    if isinstance(item, list):
        for subitem in item:
            yield from process_toc_item(subitem, level)
    elif isinstance(item, dict):
        title = item.get('/Title', '')
        yield (level, title)
        if '/First' in item:
            yield from process_toc_item(item['/First'], level + 1)
        if '/Next' in item:
            yield from process_toc_item(item['/Next'], level)
    elif hasattr(item, 'title'):  # pypdf の新しいバージョン用
        yield (level, item.title)
        if hasattr(item, 'children'):
            for child in item.children:
                yield from process_toc_item(child, level + 1)

def create_docx_with_toc(toc, content, output_path):
    doc = Document()

    # スタイルの設定
    styles = doc.styles
    for i in range(1, 4):  # TOC1, TOC2, TOC3 スタイルを作成
        style_name = f'TOC {i}'
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(14 - i)  # レベルに応じてフォントサイズを調整

    # 目次の作成
    doc.add_heading('Table of Contents', level=0)
    for level, title in process_toc_item(toc):
        doc.add_paragraph(title, style=f'TOC {min(level + 1, 3)}')

    # 本文の追加
    doc.add_page_break()
    doc.add_paragraph(content)

    doc.save(output_path)

# メイン処理
pdf_path = 'input.pdf'
docx_path = 'output.docx'

toc, content = extract_toc_and_content(pdf_path)
create_docx_with_toc(toc, content, docx_path)
